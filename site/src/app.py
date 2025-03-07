from flask import Flask, render_template, request, redirect, url_for, session, send_file
from flask_sqlalchemy import SQLAlchemy
import os
import json
from docx import Document

app = Flask(__name__)
app.debug = True
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///exercices.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'your_secret_key'  # Nécessaire pour utiliser flash messages
db = SQLAlchemy(app)

from models import Exercice

@app.route('/')
def home():
    exercices = Exercice.query.all()
    selected_exercises = session.get('selected_exercises', [])
    return render_template('index.html', exercices=exercices, selected_exercises=selected_exercises)


@app.route('/edit_exercise', methods=['GET', 'POST'])
def edit_exercise():
    if request.method == 'POST':
        # Traitement du formulaire
        exercice_id = request.form.get('id')
        exercice = Exercice.query.get(exercice_id)
        if exercice is None:
            print('Exercice non trouvé.')
            return redirect(url_for('home'))

        exercice.level = request.form.get('level')
        exercice.theme = request.form.get('theme')
        exercice.content = request.form.get('content')
        exercice.latex_code = request.form.get('latex_code')
        exercice.correction = request.form.get('correction')
        exercice.latex_correction = request.form.get('latex_correction')
        db.session.commit()
        return redirect(url_for('home'))
    else:
        exercice_id = request.args.get('id')
        exercice = Exercice.query.get(exercice_id)
        if exercice is None:
            return redirect(url_for('home'))
        return render_template('edit_exercise.html', exercice=exercice)


@app.route('/select_exercise', methods=['POST'])
def select_exercise():
    exercice_id = request.form.get('id')
    print('Exercice sélectionné:', exercice_id)
    selected_exercises = session.get('selected_exercises', [])
    if exercice_id in selected_exercises:
        selected_exercises.remove(exercice_id)
    else:
        selected_exercises.append(exercice_id)
    session['selected_exercises'] = selected_exercises

    print('Exercices sélectionnés:', selected_exercises)
    return redirect(url_for('home'))


@app.route('/export_menu')
def export_menu():
    selected_exercises = session.get('selected_exercises', [])
    if not selected_exercises:
        return "Aucun exercice sélectionné", 400

    exercices = Exercice.query.filter(Exercice.id.in_(selected_exercises)).all()
    return render_template('export.html', exercices=exercices)

@app.route('/export_exercises', methods=['GET'])
def export_exercises():
    selected_exercises = session.get('selected_exercises', [])
    if not selected_exercises:
        return "Aucun exercice sélectionné", 400

    exercices = Exercice.query.filter(Exercice.id.in_(selected_exercises)).all()
    exercices_data = [
        {
            'id': exercice.id,
            'level': exercice.level,
            'theme': exercice.theme,
            'content': exercice.content,
            'latex_code': exercice.latex_code,
            'correction': exercice.correction,
            'latex_correction': exercice.latex_correction
        }
        for exercice in exercices
    ]

    export_format = request.args.get('format', 'json')
    if export_format == 'tex':
        export_file_path = '/tmp/selected_exercises.tex'
        with open(export_file_path, 'w') as export_file:
            for exercice in exercices_data:
                export_file.write(f"\\section*{{{exercice['level']} - {exercice['theme']}}}\n")
                export_file.write(f"{exercice['content']}\n\n")
                export_file.write(f"\\subsection*{{Correction}}\n")
                export_file.write(f"{exercice['correction']}\n\n")
    elif export_format == 'docx':
        export_file_path = '/tmp/selected_exercises.docx'
        document = Document()
        for exercice in exercices_data:
            document.add_heading(f"{exercice['level']} - {exercice['theme']}", level=1)
            document.add_paragraph(exercice['content'])
            document.add_heading("Correction", level=2)
            document.add_paragraph(exercice['correction'])
        document.save(export_file_path)
    else:
        export_file_path = '/tmp/selected_exercises.json'
        with open(export_file_path, 'w') as export_file:
            json.dump(exercices_data, export_file)

    return send_file(export_file_path, as_attachment=True, download_name=f'selected_exercises.{export_format}')

if __name__ == '__main__':
    app.run(debug=True)