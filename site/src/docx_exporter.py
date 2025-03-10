from docx import Document
from htmldocx import HtmlToDocx

# do stuff to document

class DocxExporter:
    def __init__(self, exercices_data):
        self.exercices_data = exercices_data
        self.parser = HtmlToDocx()

    def export(self, file_path):
        doc = Document()
        for exercice in self.exercices_data:
            doc.add_heading(f"{exercice['level']} - {exercice['theme']}", level=1)


            self.parser.add_html_to_document(exercice['content'], doc)
            
            #doc.add_paragraph(exercice['content'])
            doc.add_heading('Correction', level=2)
            self.parser.add_html_to_document(exercice['correction'], doc)
            #doc.add_paragraph()
        doc.save(file_path)